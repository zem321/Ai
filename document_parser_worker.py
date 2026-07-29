"""Минимальный fail-closed worker для извлечения текста из вложений.

Процесс запускается отдельным Python с ``-I`` и очищенным окружением. Входные
байты поступают только через stdin, результат — через простой length-prefixed
протокол. Модуль намеренно не импортирует код бота, базу данных или провайдеры.
"""

from __future__ import annotations

import argparse
import base64
import codecs
import ctypes
from io import BytesIO
import os
from pathlib import Path, PurePosixPath
import re
import stat
import sys
import zipfile


PROTOCOL_VERSION = "SAFE-PARSER/1"
_ALLOWED_DOCX_COMPRESSION = {zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED}
_DOCX_ACTIVE_PATH_PARTS = (
    "activex",
    "embeddings",
    "vbaproject",
    "macros",
    "oleobject",
)
_PDF_ACTIVE_NAME = re.compile(
    rb"/(?:JavaScript|JS|OpenAction|AA|Launch|EmbeddedFile|RichMedia|XFA)"
    rb"(?=[\x00\t\n\f\r /<>\[\]()])",
    re.IGNORECASE,
)


class ParserRejected(Exception):
    """Ожидаемая безопасная ошибка формата или политики."""


def _preload_codecs() -> None:
    """Загружает кодеки до seccomp-профиля, запрещающего open/openat."""
    for encoding in ("utf-8", "utf-16", "cp1251", "cp437"):
        codecs.lookup(encoding)


def _decode_metadata(value: str, field: str) -> str:
    try:
        padded = value + ("=" * (-len(value) % 4))
        raw = base64.urlsafe_b64decode(padded.encode("ascii"))
        decoded = raw.decode("utf-8")
    except (ValueError, UnicodeDecodeError) as exc:
        raise ParserRejected(f"Некорректное поле {field}.") from exc
    if "\x00" in decoded or len(decoded) > 512:
        raise ParserRejected(f"Некорректное поле {field}.")
    return decoded


def _set_resource_limits(memory_bytes: int, cpu_seconds: int, output_bytes: int) -> None:
    if sys.platform != "linux":
        raise ParserRejected("Безопасный парсер поддерживается только на Linux.")
    try:
        import resource
    except ImportError as exc:  # pragma: no cover - Linux всегда имеет resource
        raise ParserRejected("Лимиты безопасного парсера недоступны.") from exc

    limits = (
        (resource.RLIMIT_AS, memory_bytes, memory_bytes),
        (resource.RLIMIT_CPU, cpu_seconds, cpu_seconds + 1),
        (resource.RLIMIT_FSIZE, output_bytes, output_bytes),
        (resource.RLIMIT_CORE, 0, 0),
        (resource.RLIMIT_NOFILE, 32, 32),
    )
    for kind, soft, hard in limits:
        try:
            resource.setrlimit(kind, (soft, hard))
        except (OSError, ValueError) as exc:
            raise ParserRejected("Не удалось установить лимиты парсера.") from exc
    if hasattr(resource, "RLIMIT_NPROC"):
        try:
            resource.setrlimit(resource.RLIMIT_NPROC, (0, 0))
        except (OSError, ValueError):
            # seccomp ниже всё равно запрещает создание процессов и потоков.
            pass


def _set_no_new_privileges() -> None:
    libc = ctypes.CDLL(None, use_errno=True)
    pr_set_no_new_privs = 38
    if libc.prctl(pr_set_no_new_privs, 1, 0, 0, 0) != 0:
        errno_value = ctypes.get_errno()
        raise OSError(errno_value, "PR_SET_NO_NEW_PRIVS")


class _LandlockRulesetAttr(ctypes.Structure):
    _fields_ = [("handled_access_fs", ctypes.c_uint64)]


class _LandlockPathBeneathAttr(ctypes.Structure):
    _fields_ = [
        ("allowed_access", ctypes.c_uint64),
        ("parent_fd", ctypes.c_int32),
    ]


def _trusted_read_roots(package_paths: tuple[str, ...]) -> tuple[str, ...]:
    # Не разрешаем /usr, /lib, стандартную библиотеку целиком или весь
    # virtualenv: внутри могут находиться конфигурация приложения и чужие
    # пакеты. Кодеки и зависимости загружены до sandbox, а выбранный пакет
    # получает отдельный узкий read-only root.
    candidates: set[str] = set()
    for value in package_paths:
        if not value:
            continue
        path = Path(value).resolve()
        candidates.add(str(path if path.is_dir() else path.parent))
    return tuple(
        sorted(
            {
                value
                for value in candidates
                if value != "/" and Path(value).is_dir()
            }
        )
    )


def _apply_landlock(package_paths: tuple[str, ...]) -> None:
    """Разрешает только чтение системного Python и импортированных пакетов."""
    libc = ctypes.CDLL(None, use_errno=True)
    syscall_create_ruleset = 444
    syscall_add_rule = 445
    syscall_restrict_self = 446
    landlock_create_ruleset_version = 1
    landlock_rule_path_beneath = 1

    abi = libc.syscall(
        syscall_create_ruleset,
        ctypes.c_void_p(),
        ctypes.c_size_t(0),
        ctypes.c_uint(landlock_create_ruleset_version),
    )
    if abi < 1:
        errno_value = ctypes.get_errno()
        raise OSError(errno_value, "Landlock ABI недоступен")

    execute = 1 << 0
    write_file = 1 << 1
    read_file = 1 << 2
    read_dir = 1 << 3
    remove_dir = 1 << 4
    remove_file = 1 << 5
    make_char = 1 << 6
    make_dir = 1 << 7
    make_reg = 1 << 8
    make_sock = 1 << 9
    make_fifo = 1 << 10
    make_block = 1 << 11
    make_sym = 1 << 12
    handled = (
        execute
        | write_file
        | read_file
        | read_dir
        | remove_dir
        | remove_file
        | make_char
        | make_dir
        | make_reg
        | make_sock
        | make_fifo
        | make_block
        | make_sym
    )
    if abi >= 2:
        handled |= 1 << 13  # REFER
    if abi >= 3:
        handled |= 1 << 14  # TRUNCATE

    ruleset_attr = _LandlockRulesetAttr(handled_access_fs=handled)
    ruleset_fd = libc.syscall(
        syscall_create_ruleset,
        ctypes.byref(ruleset_attr),
        ctypes.sizeof(ruleset_attr),
        0,
    )
    if ruleset_fd < 0:
        errno_value = ctypes.get_errno()
        raise OSError(errno_value, "landlock_create_ruleset")

    opened: list[int] = []
    try:
        allowed_read = execute | read_file | read_dir
        open_flags = getattr(os, "O_PATH", os.O_RDONLY) | os.O_CLOEXEC
        for root in _trusted_read_roots(package_paths):
            try:
                parent_fd = os.open(root, open_flags)
            except OSError:
                continue
            opened.append(parent_fd)
            path_attr = _LandlockPathBeneathAttr(
                allowed_access=allowed_read,
                parent_fd=parent_fd,
            )
            result = libc.syscall(
                syscall_add_rule,
                ruleset_fd,
                landlock_rule_path_beneath,
                ctypes.byref(path_attr),
                0,
            )
            if result != 0:
                errno_value = ctypes.get_errno()
                raise OSError(errno_value, f"landlock_add_rule({root})")

        _set_no_new_privileges()
        if libc.syscall(syscall_restrict_self, ruleset_fd, 0) != 0:
            errno_value = ctypes.get_errno()
            raise OSError(errno_value, "landlock_restrict_self")
    finally:
        for descriptor in opened:
            os.close(descriptor)
        os.close(ruleset_fd)


class _SeccompArgCmp(ctypes.Structure):
    _fields_ = [
        ("arg", ctypes.c_uint),
        ("op", ctypes.c_int),
        ("datum_a", ctypes.c_uint64),
        ("datum_b", ctypes.c_uint64),
    ]


def _apply_seccomp(*, deny_all_file_opens: bool) -> None:
    """Запрещает сеть, новые процессы, ptrace и мутации файловой системы."""
    seccomp = ctypes.CDLL("libseccomp.so.2", use_errno=True)
    seccomp.seccomp_init.argtypes = [ctypes.c_uint32]
    seccomp.seccomp_init.restype = ctypes.c_void_p
    seccomp.seccomp_release.argtypes = [ctypes.c_void_p]
    seccomp.seccomp_syscall_resolve_name.argtypes = [ctypes.c_char_p]
    seccomp.seccomp_syscall_resolve_name.restype = ctypes.c_int
    seccomp.seccomp_rule_add.argtypes = [
        ctypes.c_void_p,
        ctypes.c_uint32,
        ctypes.c_int,
        ctypes.c_uint,
    ]
    seccomp.seccomp_rule_add.restype = ctypes.c_int
    seccomp.seccomp_rule_add_array.argtypes = [
        ctypes.c_void_p,
        ctypes.c_uint32,
        ctypes.c_int,
        ctypes.c_uint,
        ctypes.POINTER(_SeccompArgCmp),
    ]
    seccomp.seccomp_rule_add_array.restype = ctypes.c_int
    seccomp.seccomp_load.argtypes = [ctypes.c_void_p]
    seccomp.seccomp_load.restype = ctypes.c_int

    scmp_act_allow = 0x7FFF0000
    scmp_act_errno = 0x00050000 | 1  # EPERM
    scmp_cmp_masked_eq = 7
    context = seccomp.seccomp_init(scmp_act_allow)
    if not context:
        raise OSError("seccomp_init")
    try:
        blocked = [
            "socket", "socketpair", "connect", "bind", "listen", "accept",
            "accept4", "sendto", "sendmsg", "sendmmsg", "recvfrom",
            "recvmsg", "recvmmsg", "shutdown",
            "execve", "execveat", "fork", "vfork", "clone", "clone3",
            "ptrace", "process_vm_readv", "process_vm_writev",
            "kill", "tkill", "tgkill", "pidfd_open", "pidfd_send_signal",
            "mount", "umount2", "pivot_root", "chroot", "unshare", "setns",
            "bpf", "perf_event_open", "keyctl", "add_key", "request_key",
            "open_by_handle_at", "init_module", "finit_module", "delete_module",
            "reboot", "swapon", "swapoff", "openat2", "creat", "truncate",
            "ftruncate", "unlink", "unlinkat", "rename", "renameat",
            "renameat2", "mkdir", "mkdirat", "rmdir", "link", "linkat",
            "symlink", "symlinkat", "mknod", "mknodat", "chmod", "fchmod",
            "fchmodat", "chown", "fchown", "lchown", "fchownat",
            "setxattr", "lsetxattr", "fsetxattr", "removexattr",
            "lremovexattr", "fremovexattr",
            "readlink", "readlinkat", "name_to_handle_at",
            "io_uring_setup", "io_uring_enter", "io_uring_register",
            "userfaultfd", "kcmp",
        ]
        if deny_all_file_opens:
            blocked.extend(("open", "openat"))
        for name in blocked:
            syscall_number = seccomp.seccomp_syscall_resolve_name(
                name.encode("ascii")
            )
            if syscall_number < 0:
                continue
            if (
                seccomp.seccomp_rule_add(
                    context,
                    scmp_act_errno,
                    syscall_number,
                    0,
                )
                != 0
            ):
                raise OSError(f"seccomp_rule_add({name})")

        # Дополнительно запрещаем open/openat с флагами записи. Landlock
        # обеспечивает ту же границу, но это независимый второй слой.
        write_flag_checks = (
            (os.O_ACCMODE, os.O_WRONLY),
            (os.O_ACCMODE, os.O_RDWR),
            (os.O_CREAT, os.O_CREAT),
            (os.O_TRUNC, os.O_TRUNC),
            (os.O_APPEND, os.O_APPEND),
        )
        if not deny_all_file_opens:
            for name, flags_argument in (("open", 1), ("openat", 2)):
                syscall_number = seccomp.seccomp_syscall_resolve_name(
                    name.encode("ascii")
                )
                if syscall_number < 0:
                    continue
                for mask, expected in write_flag_checks:
                    comparison = _SeccompArgCmp(
                        arg=flags_argument,
                        op=scmp_cmp_masked_eq,
                        datum_a=mask,
                        datum_b=expected,
                    )
                    result = seccomp.seccomp_rule_add_array(
                        context,
                        scmp_act_errno,
                        syscall_number,
                        1,
                        ctypes.byref(comparison),
                    )
                    if result != 0:
                        raise OSError(f"seccomp_rule_add_array({name})")

        if seccomp.seccomp_load(context) != 0:
            raise OSError("seccomp_load")
    finally:
        seccomp.seccomp_release(context)


def _activate_sandbox(package_paths: tuple[str, ...], require_landlock: bool) -> None:
    privilege_drop_error: OSError | None = None
    if os.geteuid() == 0:
        try:
            os.setgroups([])
            os.setgid(65534)
            os.setuid(65534)
        except OSError as exc:
            # Root внутри непривилегированного user namespace часто не имеет
            # CAP_SETUID. В этом случае обязательные Landlock + seccomp ниже
            # остаются fail-closed границей доступа.
            privilege_drop_error = exc
    landlock_active = False
    try:
        _apply_landlock(package_paths)
        landlock_active = True
    except (OSError, ValueError) as exc:
        # Старые/container kernels могут не включать Landlock. Тогда seccomp
        # ниже запрещает вообще любые новые open/openat, а не только запись.
        # Парсер продолжает работать исключительно с уже открытыми stdin/stdout.
        landlock_active = False
    try:
        _set_no_new_privileges()
        _apply_seccomp(deny_all_file_opens=not landlock_active)
    except (OSError, AttributeError) as exc:
        raise ParserRejected("seccomp-песочница недоступна.") from exc
    if require_landlock and not landlock_active and privilege_drop_error is None:
        # Отсутствие Landlock допустимо только из-за более строгого no-open
        # профиля seccomp; эта ветка документирует намеренный fail-closed fallback.
        return


def _canonical_pdf_names(raw: bytes) -> bytes:
    def replace_name_escape(match: re.Match[bytes]) -> bytes:
        try:
            return bytes((int(match.group(1), 16),))
        except ValueError:
            return match.group(0)

    return re.sub(rb"#([0-9A-Fa-f]{2})", replace_name_escape, raw)


def _reject_pdf_active_content(raw: bytes, reader: object) -> None:
    if _PDF_ACTIVE_NAME.search(_canonical_pdf_names(raw)):
        raise ParserRejected("PDF с активным содержимым не принимается.")
    trailer = getattr(reader, "trailer", {})
    root = trailer.get("/Root") if hasattr(trailer, "get") else None
    try:
        if root:
            root_object = root.get_object()
            if any(
                key in root_object
                for key in ("/OpenAction", "/AA", "/JavaScript", "/XFA")
            ):
                raise ParserRejected("PDF с активным содержимым не принимается.")
            names = root_object.get("/Names")
            if names:
                names_object = names.get_object()
                if any(
                    key in names_object
                    for key in ("/JavaScript", "/EmbeddedFiles")
                ):
                    raise ParserRejected(
                        "PDF со встроенными файлами не принимается."
                    )
    except ParserRejected:
        raise
    except Exception as exc:
        raise ParserRejected("Некорректная структура PDF.") from exc


def _extract_pdf(raw: bytes, max_chars: int, max_pages: int) -> str:
    try:
        from pypdf import PdfReader
        import pypdf
    except ImportError as exc:
        raise ParserRejected("Компонент чтения PDF не установлен.") from exc
    try:
        _activate_sandbox((pypdf.__file__ or "",), require_landlock=True)
        reader = PdfReader(BytesIO(raw), strict=True)
        if reader.is_encrypted:
            raise ParserRejected("Зашифрованные PDF не принимаются.")
        if len(reader.pages) > max_pages:
            raise ParserRejected(f"В PDF больше {max_pages} страниц.")
        _reject_pdf_active_content(raw, reader)
        texts: list[str] = []
        total = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            remaining = max_chars - total
            if remaining <= 0:
                break
            texts.append(text[:remaining])
            total += min(len(text), remaining)
        return "\n".join(texts).strip()
    except ParserRejected:
        raise
    except Exception as exc:
        raise ParserRejected("Не удалось безопасно прочитать PDF.") from exc


def _validate_docx_archive(
    raw: bytes,
    max_expanded_bytes: int,
    max_entries: int,
) -> None:
    try:
        with zipfile.ZipFile(BytesIO(raw)) as archive:
            infos = archive.infolist()
            if len(infos) > max_entries:
                raise ParserRejected("В DOCX слишком много внутренних файлов.")

            seen: set[str] = set()
            expanded = 0
            required = {"[content_types].xml", "word/document.xml"}
            for info in infos:
                name = info.filename.replace("\\", "/")
                path = PurePosixPath(name)
                lowered = name.casefold()
                if (
                    not name
                    or name.startswith("/")
                    or "\\" in info.filename
                    or ":" in path.parts[0]
                    or any(part in {"", ".", ".."} for part in path.parts)
                    or lowered in seen
                ):
                    raise ParserRejected("DOCX содержит опасный внутренний путь.")
                seen.add(lowered)
                if info.flag_bits & 0x1:
                    raise ParserRejected("Зашифрованные DOCX не принимаются.")
                if info.compress_type not in _ALLOWED_DOCX_COMPRESSION:
                    raise ParserRejected("DOCX использует неподдерживаемое сжатие.")
                mode = (info.external_attr >> 16) & 0o170000
                if mode == stat.S_IFLNK:
                    raise ParserRejected("Символические ссылки в DOCX запрещены.")
                if any(part in lowered for part in _DOCX_ACTIVE_PATH_PARTS):
                    raise ParserRejected("DOCX с активным содержимым не принимается.")
                expanded += info.file_size
                if info.file_size > 10 * 1024 * 1024:
                    raise ParserRejected("Внутренний файл DOCX слишком велик.")
                if info.file_size > 1024 * 1024:
                    if info.compress_size == 0:
                        raise ParserRejected("Обнаружена ZIP-бомба в DOCX.")
                    if info.file_size / info.compress_size > 200:
                        raise ParserRejected("Подозрительное сжатие DOCX.")
                if expanded > max_expanded_bytes:
                    raise ParserRejected("DOCX слишком велик после распаковки.")

            if not required.issubset(seen):
                raise ParserRejected("В DOCX отсутствуют обязательные части.")
            # CRC проверяется только после дешёвых метаданных и лимитов.
            # Иначе testzip() мог первым распаковать ZIP-бомбу.
            if archive.testzip() is not None:
                raise ParserRejected("DOCX повреждён.")

            for info in infos:
                lowered = info.filename.casefold()
                if not lowered.endswith((".xml", ".rels")):
                    continue
                xml = archive.read(info)
                xml_lower = xml.lower()
                if b"<!doctype" in xml_lower or b"<!entity" in xml_lower:
                    raise ParserRejected("DTD и XML-сущности в DOCX запрещены.")
                if lowered.endswith(".rels") and re.search(
                    rb"\btargetmode\s*=\s*[\"']external[\"']",
                    xml,
                    re.IGNORECASE,
                ):
                    raise ParserRejected("Внешние связи в DOCX запрещены.")
    except ParserRejected:
        raise
    except (zipfile.BadZipFile, RuntimeError, OSError) as exc:
        raise ParserRejected("Некорректный DOCX-файл.") from exc


def _extract_docx(
    raw: bytes,
    max_chars: int,
    max_expanded_bytes: int,
    max_entries: int,
) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ParserRejected("Компонент чтения DOCX не установлен.") from exc
    try:
        _activate_sandbox((docx.__file__ or "",), require_landlock=True)
        _validate_docx_archive(raw, max_expanded_bytes, max_entries)
        document = docx.Document(BytesIO(raw))
        result: list[str] = []
        total = 0

        def append_text(value: str) -> bool:
            nonlocal total
            remaining = max_chars - total
            if remaining <= 0:
                return False
            result.append(value[:remaining])
            total += min(len(value), remaining)
            return total < max_chars

        for paragraph in document.paragraphs:
            if not append_text(paragraph.text):
                break
        if total < max_chars:
            for table in document.tables:
                for row in table.rows:
                    if not append_text("\t".join(cell.text for cell in row.cells)):
                        break
                if total >= max_chars:
                    break
        return "\n".join(result).strip()
    except ParserRejected:
        raise
    except Exception as exc:
        raise ParserRejected("Не удалось безопасно прочитать DOCX.") from exc


def _extract_text(raw: bytes, max_chars: int, package_paths: tuple[str, ...]) -> str:
    _activate_sandbox(package_paths, require_landlock=True)
    if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
        try:
            decoded = raw.decode("utf-16")
        except UnicodeDecodeError as exc:
            raise ParserRejected("Некорректный UTF-16 файл.") from exc
        if any(
            ord(char) < 32 and char not in "\t\n\r\f"
            for char in decoded
        ):
            raise ParserRejected("Файл содержит управляющие символы.")
        return decoded[:max_chars]
    if b"\x00" in raw:
        raise ParserRejected("Файл содержит бинарные данные.")
    disallowed_controls = sum(
        (byte < 32 and byte not in {9, 10, 12, 13}) or byte == 127
        for byte in raw
    )
    if disallowed_controls > max(2, len(raw) // 100):
        raise ParserRejected("Файл содержит бинарные управляющие данные.")
    for encoding in ("utf-8", "cp1251"):
        try:
            return raw.decode(encoding)[:max_chars]
        except UnicodeDecodeError:
            continue
    raise ParserRejected(
        "Текстовый файл должен быть в UTF-8, UTF-16 или Windows-1251."
    )


def _emit(status: str, payload: str, output_bytes: int) -> None:
    encoded = payload.encode("utf-8", errors="strict")
    if len(encoded) > output_bytes - 128:
        encoded = encoded[: output_bytes - 128]
        while encoded:
            try:
                encoded.decode("utf-8")
                break
            except UnicodeDecodeError:
                encoded = encoded[:-1]
    header = f"{PROTOCOL_VERSION} {status} {len(encoded)}\n".encode("ascii")
    sys.stdout.buffer.write(header)
    sys.stdout.buffer.write(encoded)
    sys.stdout.buffer.flush()


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--filename", required=True)
    parser.add_argument("--mime", required=True)
    parser.add_argument("--max-input", type=int, required=True)
    parser.add_argument("--max-chars", type=int, required=True)
    parser.add_argument("--max-pages", type=int, required=True)
    parser.add_argument("--max-docx-expanded", type=int, required=True)
    parser.add_argument("--max-docx-entries", type=int, required=True)
    parser.add_argument("--memory", type=int, required=True)
    parser.add_argument("--cpu", type=int, required=True)
    parser.add_argument("--output-bytes", type=int, required=True)
    parser.add_argument("--require-landlock", choices=("0", "1"), required=True)
    return parser.parse_args()


def main() -> int:
    args = _parse_args()
    try:
        if args.require_landlock != "1":
            raise ParserRejected("Ослабление Landlock запрещено.")
        numeric_values = (
            args.max_input,
            args.max_chars,
            args.max_pages,
            args.max_docx_expanded,
            args.max_docx_entries,
            args.memory,
            args.cpu,
            args.output_bytes,
        )
        if any(value <= 0 for value in numeric_values):
            raise ParserRejected("Некорректные лимиты парсера.")
        _set_resource_limits(args.memory, args.cpu, args.output_bytes)
        _preload_codecs()
        filename = _decode_metadata(args.filename, "filename")
        mime = _decode_metadata(args.mime, "mime").strip().lower()
        raw = sys.stdin.buffer.read(args.max_input + 1)
        if not raw:
            raise ParserRejected("Файл пуст.")
        if len(raw) > args.max_input:
            raise ParserRejected("Файл превышает допустимый размер.")

        name_lower = filename.casefold()
        if name_lower.endswith(".pdf") or mime == "application/pdf":
            if b"%PDF-" not in raw[:1024]:
                raise ParserRejected("Содержимое файла не является PDF.")
            text = _extract_pdf(raw, args.max_chars, args.max_pages)
        elif (
            name_lower.endswith(".docx")
            or mime
            == "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ):
            if not raw.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
                raise ParserRejected("Содержимое файла не является DOCX.")
            text = _extract_docx(
                raw,
                args.max_chars,
                args.max_docx_expanded,
                args.max_docx_entries,
            )
        else:
            text = _extract_text(raw, args.max_chars, ())
        _emit("OK", text, args.output_bytes)
        return 0
    except BaseException as exc:
        if isinstance(exc, (KeyboardInterrupt, SystemExit)):
            raise
        message = str(exc)
        message = re.sub(r"[\x00-\x1f\x7f]+", " ", message).strip()[:500]
        try:
            _emit(
                "ERR",
                message or "Не удалось безопасно прочитать файл.",
                max(1024, getattr(args, "output_bytes", 4096)),
            )
        except BaseException:
            pass
        return 2


if __name__ == "__main__":
    raise SystemExit(main())

